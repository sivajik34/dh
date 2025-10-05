# ============================================================================
# KNOWLEDGE BASE INGESTION SERVICE (knowledge_ingestion/main.py)
# ============================================================================
from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from typing import List
import PyPDF2
import docx
from bs4 import BeautifulSoup
import io
from vector_store import VectorStore
app = FastAPI(title="Knowledge Base Ingestion")

vector_store = VectorStore()

class Document(BaseModel):
    id: str
    title: str
    content: str
    category: str
    metadata: dict = {}

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF"""
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX"""
    doc = docx.Document(io.BytesIO(file_content))
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_html(file_content: bytes) -> str:
    """Extract text from HTML"""
    soup = BeautifulSoup(file_content, 'html.parser')
    return soup.get_text()

@app.post("/ingest/document")
async def ingest_document(doc: Document):
    """Ingest a single document"""
    vector_store.add_documents([doc.dict()])
    return {"status": "success", "document_id": doc.id}

@app.post("/ingest/file")
async def ingest_file(file: UploadFile = File(...), category: str = "general"):
    """Ingest a file (PDF, DOCX, HTML, TXT)"""
    content = await file.read()
    
    # Extract text based on file type
    if file.filename.endswith('.pdf'):
        text = extract_text_from_pdf(content)
    elif file.filename.endswith('.docx'):
        text = extract_text_from_docx(content)
    elif file.filename.endswith('.html'):
        text = extract_text_from_html(content)
    else:
        text = content.decode('utf-8')
    
    # Create document
    doc = {
        'id': file.filename,
        'title': file.filename,
        'content': text,
        'category': category,
        'metadata': {'filename': file.filename}
    }
    
    vector_store.add_documents([doc])
    
    return {"status": "success", "filename": file.filename, "length": len(text)}

@app.post("/search")
async def search_knowledge(query: str, k: int = 5):
    """Search knowledge base"""
    results = vector_store.search(query, k)
    return {"query": query, "results": results}

@app.delete("/document/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document"""
    vector_store.delete_by_id([doc_id])
    return {"status": "success", "deleted": doc_id}

@app.get("/health")
async def health():
    return {"status": "healthy", "service": "knowledge_ingestion"}